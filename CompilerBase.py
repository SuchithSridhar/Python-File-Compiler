"""
This is a program for the base of
the program for a GUI this program
is not runnable as there is nothing
but functions here - 2 programs
GUI and Console Base Program

functions :
tab_space(file_name)
    # Converts tabs to spaces
output(file, output_data)
    : Returns Output data of a perticular file/program
create_output(questions_list)
    # Creates temp.py and runs it for the outputs of the Qs
copy_clip()
    : Returns the output copied from the clipboard
write_file(values_dict, file_to_write="Compiled Files.txt")
    : Returns file_name after writing the Output into Notepad
delete_temp(temp_path="temp.py")
    : Returns True if successful

the_main(questions_list, values_dict, wait_func, file_to_write="Compiled Files.txt") :
Does everything needed :)
example vals :
vals_dict = {
    'title': "list based assignment",
    'by_line': "suchith sridhar",
    'files': q_list,
    'question_line': True,
    'doc_string': "Name : Suchith\nClass:12 C",
    'output': True,
    'output_data': None,
    'output_start': string
    'output_end': string
}
"""

import os
import pyperclip
import time
import docx
from docx.shared import Pt
from docx.shared import Inches


def tab_space(file_name):
    """
        Function to replace all tabs in a file
        with 4 spaces for better printing
        """

    with open(file_name, 'r') as file:
        data = file.read()
    data = data.replace('\t', '    ')
    with open(file_name, 'w') as file:
        file.write(data)


def output(file, output_data):
    """
        This function return the output of the file
        that has been provided as the file parameter
        output_data is the output of running all the programs
        """
    finding = "---#$" + file + "#$---"
    start = output_data.find(finding) + len(finding)
    end = output_data.find("---#$", start)

    if end == -1:
        end = len(output_data)

    # Returns the output of the respective file
    return output_data[start:end]


def create_output(questions_list):
    """
        A function that create a new python program that runs all
        the programs from file temp.py
        """
    with open("temp.py", 'w') as file:
        file.write("#Temporary file to be deleted at all costs\n")
        for ques in questions_list:
            file.write("print('" + ("---#$" + ques + "#$---") + "')" + '\n')
            file.write("try:\n")
            file.write("\timport " + ques.replace('.py', '') + '\n')
            file.write("except:\n")
            file.write("\tprint(' --------------- ERROR IN FILE :" +
                       ques + " ---------------')" + '\n')

        file.write("print('\\n\\nPRESS CTRL+A then CTRL+C then ENTER')\n")
        file.write("input()\n")
        string = "f=open('asdvgfrehtgdv.txt', 'w');f.close()"
        file.write(string)

    os.startfile(os.path.join(os.getcwd(), 'temp.py'))


def copy_clip():
    output_copied = pyperclip.paste()
    output_copied = output_copied.replace(
        "PRESS CTRL+A then CTRL+C then ENTER", "")
    output_copied = output_copied[:-3]
    return output_copied


def write_docx(values_dict, file_to_write="Compiled Files.docx"):
    my_docx = docx.Document()
    style = my_docx.styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt(12)

    paragraph_format = my_docx.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    paragraph_format.line_spacing = 1

    margin = 1

    sections = my_docx.sections
    for section in sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    string = ""
    top_string = ""

    if values_dict['title'] not in (None, False):
        top_string += "\t\t" + values_dict['title'] + '\n'

    if values_dict['by_line'] not in (None, False):
        top_string += "\t\t" + values_dict["by_line"] + '\n'

    my_docx.add_paragraph(top_string)

    for file in values_dict['files']:
        # Files has to be a list of all the questions
        string = ""

        if values_dict['question_line'] not in (None, False):
            string += file.replace('.py', '') + '-' * 50
            string += "\n"

        # doc_string has to be a Doc string
        # without the doc string quotes
        if values_dict['doc_string'] not in (None, False):
            string += "'''\n"
            str_var = file.replace('.py', '')
            string += ("Program: "+str_var + ("\n") +
                       (values_dict['doc_string']) + ("\n'''\n"))

        with open(file, 'r') as read_q:
            data = read_q.read()
            string += (data)

        if values_dict["output"] not in (None, False):
            if values_dict["output_start"] not in (None, False):
                string += (values_dict["output_start"])

            str_var = (output(file, values_dict["output_data"]))

            if "\r" in str_var:
                str_var = str_var.replace('\r', '')

            # ------------------- This is IMPORTANT -------------------
            # We replace all the \n because in the computer console version
            # there are actually 2 escape characters \n and \r so when we write
            # it without removing the \n there is 2 spaces not one
            # due to this problem
            # https://en.wikipedia.org/wiki/Newline

            # Example :
            # string = string.replace('\n', '#')
            # string = string.replace('\r', '%')
            # %#Enter number of elements in list: 5%#0.79  0.92  0.62
            # 0.89 0.09%#sum = 3.31 avg = 0.66%#%#Press Enter to Exit%#%#%

            string += str_var
            if values_dict["output_end"] not in (None, False):
                string += values_dict["output_end"]

            string += "\n\n\n"

            my_docx.add_paragraph(string)
            my_docx.add_page_break()

    my_docx.save(file_to_write)
    return file_to_write


def write_file(values_dict, file_to_write="Compiled Files.txt"):
    """
        This function writes the final file from which the text
        is to be copied.
        The values_dict include :

        + title : at the top
        + by_line : has to be just the name
        + files : list of all the files that have to
                be read
        + question_line : anything but none and false will do
        + doc_sting_start : The first few lines of the code thats
                required by sir
        + output : anything but none and false will do
        + output_data : output_data of all the files
                this data can be collected through the function
                create_output in the compiler_base.py file
        + output_start: string
    + output_end: string
        """

    with open(file_to_write, 'w') as main_file:
        if values_dict['title'] not in (None, False):
            string = "\t\t" + values_dict['title'] + '\n'
            main_file.write(string)

        if values_dict['by_line'] not in (None, False):
            string = "\t\t" + values_dict["by_line"] + '\n'
            main_file.write(string)

        for file in values_dict['files']:
            # Files has to be a list of all the questions

            if values_dict['question_line'] not in (None, False):
                string = file.replace('.py', '') + '-' * 50
                main_file.write(string)
                main_file.write("\n")

            # doc_string has to be a Doc string
            # without the doc string quotes
            if values_dict['doc_string'] not in (None, False):
                main_file.write("'''\n")
                string = string.replace('-' * 50, '')
                main_file.write("Program: "+string)
                main_file.write("\n")
                main_file.write(values_dict['doc_string'])
                main_file.write("\n'''\n")

            with open(file, 'r') as read_q:
                data = read_q.read()
                main_file.write(data)

            if values_dict["output"] not in (None, False):
                if values_dict["output_start"] not in (None, False):
                    main_file.write(values_dict["output_start"])

                string = (output(file, values_dict["output_data"]))

                if "\r" in string:
                    string = string.replace('\r', '')

                main_file.write(string)
                if values_dict["output_end"] not in (None, False):
                    main_file.write(values_dict["output_end"])

                main_file.write("\n\n\n")

    return file_to_write


def delete_temp(temp_path="temp.py"):
    """
        Function returns True is succesfully deleted temp.py
        else returns false
        temp arg to be provided if the file present
        in a different folded
        """
    try:
        os.remove(temp_path)
        return True
    except:
        return False


def the_main(questions_list, values_dict,
             wait_func=lambda: 1, file_to_write="Compiled Files.txt"):
    # IMPORTANT NOTE !!
    # The questions and the temp.py file has to be present in the same
    # folder so that temp.py can import them.
    if values_dict["output"] not in (None, False):
        create_output(questions_list)
        # Creates a file from where all the programs are run

        while not os.path.exists('asdvgfrehtgdv.txt'):
            time.sleep(1)
            pass

        try:
            os.remove('asdvgfrehtgdv.txt')

        except FileNotFoundError:
            pass
        # waits for an event so that the user can confirm that

        output_data = copy_clip()
        # Copy the output from the clipborad

        delete_temp()
        print("Got here - basic 1")
        # delete the temp.py file if possible, provide a path if the file
        # not present in the cwd

        values_dict["output_data"] = output_data

    else:
        values_dict["output_data"] = None

    if values_dict["word_file"]:
        if ".txt" in file_to_write:
            file_to_write = "Compiled Files.docx"

        print("Got here - basic 2")
        file_written = write_docx(values_dict, file_to_write)
        print("Got here - basic 3")

    else:
        file_written = write_file(values_dict, file_to_write)
        # Writes the file into a notepad file so that it can be written there

        # Changes all tabs to spaces for better printing
        tab_space(file_written)

    os.startfile(file_written)


if __name__ == "__main__":
    print("Just a sample :")
    q_list = []
    while True:
        chs = input("Enter questions :")
        if chs.lower() in ('stop', ""):
            break
        else:
            try:
                open(chs)
            except FileNotFoundError:
                print("error in file opening")
                continue
            else:
                q_list.append(chs)

    vals_dict = {
        'title': "      list based assignment",
        'by_line': "        - suchith sridhar",
        'files': q_list,
        'question_line': True,
        'doc_string': "Name : Suchith\nClass:12 C",
        'output': True,
        'output_data': None,
        'output_start': "---- Output ----",
        "output_end": "",
        "word_file": False
    }

    the_main(q_list, vals_dict)

    print("done")
    input()
